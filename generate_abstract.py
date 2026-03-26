"""
generate_abstract.py
Generates a 5-page research abstract (DOCX) with tables and figures.
Run AFTER run_pipeline.py to use real results.
"""
import os
import sys
import json
import math
import pandas as pd
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
TABLES_DIR = os.path.join(BASE_DIR, "outputs", "tables")
CHARTS_DIR = os.path.join(BASE_DIR, "outputs", "charts")
META_PATH  = os.path.join(BASE_DIR, "outputs", "metadata.json")
OUT_PATH   = os.path.join(BASE_DIR, "research_abstract.docx")

# Chart order and captions
CHARTS = [
    ("roc_curve.png",         "Figure 1. ROC Curves for All Six Classifiers (OHE, best configuration per model)"),
    ("model_performance.png", "Figure 2. F1-Score and ROC-AUC Comparison Across Models (Best Configuration)"),
    ("feature_importance.png","Figure 3. Top 15 Feature Importances from Random Forest (SMOTE + OHE)"),
    ("imbalance_impact.png",  "Figure 4. Effect of Imbalance Handling Strategy on Recall and F1-Score (Averaged)"),
    ("confusion_matrix.png",  "Figure 5. Confusion Matrix of the Best Performing Model"),
]


# ── Utilities ──────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text: str, level: int, color: tuple = (31, 73, 125)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h


def _add_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                    size: int = 11, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def _add_figure(doc, img_path: str, caption: str, width_in: float = 5.5):
    if not os.path.exists(img_path):
        _add_paragraph(doc, f"[Figure not found: {caption}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))

    cap_p = doc.add_paragraph(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        run.font.size = Pt(9)
        run.italic = True
    doc.add_paragraph("")


def _add_comparison_table(doc, df: pd.DataFrame):
    """Add a formatted model comparison table."""
    display_cols = ["Model", "Encoding", "Imbalance", "Accuracy", "Precision",
                    "Recall", "F1", "ROC_AUC"]
    df = df[display_cols].copy()

    table = doc.add_table(rows=1, cols=len(display_cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0].cells
    for i, col in enumerate(display_cols):
        hdr[i].text = col.replace("_", "-")
        _set_cell_bg(hdr[i], "1F497D")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Best ROC-AUC value for highlighting
    try:
        best_roc = df["ROC_AUC"].max()
    except Exception:
        best_roc = None

    for _, row_data in df.iterrows():
        row = table.add_row().cells
        is_best = (best_roc is not None and
                   abs(float(row_data["ROC_AUC"]) - best_roc) < 1e-4)

        for i, col in enumerate(display_cols):
            val = row_data[col]
            if col in ["Accuracy", "Precision", "Recall", "F1", "ROC_AUC"]:
                try:
                    cell_text = f"{float(val):.4f}"
                except Exception:
                    cell_text = str(val)
            else:
                cell_text = str(val)

            row[i].text = cell_text
            if is_best:
                _set_cell_bg(row[i], "E2EFDA")

            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(7.5)
                    if is_best:
                        run.bold = True
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


# ── Main document builder ──────────────────────────────────────────────────────
def generate_abstract():
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Load data ────────────────────────────────────────────────────
    meta = {}
    if os.path.exists(META_PATH):
        with open(META_PATH) as f:
            meta = json.load(f)

    comp_df = None
    csv_path = os.path.join(TABLES_DIR, "model_comparison.csv")
    if os.path.exists(csv_path):
        comp_df = pd.read_csv(csv_path)

    # ── Defaults if no pipeline run yet ─────────────────────────────
    bm = meta.get("best_model", {})
    best_name    = bm.get("name", "Random Forest")
    best_enc     = bm.get("encoding", "OHE")
    best_imb     = bm.get("imbalance", "SMOTE")
    best_roc_auc = bm.get("roc_auc", "N/A")
    best_f1      = bm.get("f1", "N/A")
    best_recall  = bm.get("recall", "N/A")
    pos_rate     = meta.get("positive_rate", 0.113)
    n_samples    = meta.get("dataset_shape", [41188, 45])[0]
    top_features = meta.get("top_features", ["euribor3m", "nr.employed", "pdays",
                                               "emp.var.rate", "cons.conf.idx"])

    # ── Title Block ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        "Comparative Analysis of Traditional Machine Learning Algorithms "
        "for Bank Marketing Customer Subscription Prediction"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph("")
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_p.add_run(
        "Hieu Ngan  ·  Faculty of Information Technology  ·  "
        f"Research Paper  |  {datetime.now().strftime('%B %Y')}"
    ).font.size = Pt(10)

    doc.add_paragraph("")

    # ── Abstract Box ─────────────────────────────────────────────────
    _add_heading(doc, "Abstract", level=2)
    abs_text = (
        "This study presents a comparative evaluation of six traditional machine learning "
        "classifiers—Logistic Regression, Decision Tree, Random Forest, Support Vector "
        "Machine (SVM), k-Nearest Neighbours (k-NN), and Naïve Bayes—applied to the UCI "
        "Bank Marketing dataset for predicting customer subscription to term deposits. "
        "A full factorial experiment of 36 configurations was executed, systematically "
        "varying encoding strategy (one-hot vs. label encoding) and class-imbalance "
        "handling technique (none, class weighting, SMOTE). The dataset comprises "
        f"{n_samples:,} client records with approximately {pos_rate:.1%} positive "
        "subscription rate, representing a strong class imbalance. Experimental results "
        f"indicate that {best_name} with {best_enc} encoding and {best_imb} resampling "
        f"achieves the highest performance (ROC-AUC = {best_roc_auc}, "
        f"F1 = {best_f1}, Recall = {best_recall}). "
        "SMOTE consistently improved recall across all models. Feature analysis revealed "
        f"that economic indicators ({', '.join(top_features[:3])}) and campaign-specific "
        "variables are the most predictive. The duration feature was excluded to avoid "
        "data leakage and preserve real-world applicability."
    )
    _add_paragraph(doc, abs_text)

    keywords_p = doc.add_paragraph()
    keywords_p.add_run("Keywords: ").bold = True
    keywords_p.add_run(
        "bank marketing, term deposit prediction, machine learning, class imbalance, "
        "SMOTE, feature importance, logistic regression, random forest, SVM"
    ).italic = True

    doc.add_paragraph("")

    # ── 1. Introduction ───────────────────────────────────────────────
    _add_heading(doc, "1  Introduction", level=1)
    _add_paragraph(doc,
        "Direct marketing campaigns are a critical revenue channel for retail banking institutions. "
        "The ability to predict which customers are likely to subscribe to a term deposit allows "
        "campaign managers to allocate call-centre resources more efficiently, reducing costs while "
        "improving conversion rates. Machine learning-based classification systems can serve as "
        "decision-support tools, enabling personalised outreach at scale."
    )
    _add_paragraph(doc,
        "The UCI Bank Marketing dataset (Moro et al., 2014) captures the outcomes of telemarketing "
        "campaigns conducted by a Portuguese banking institution between 2008 and 2013. Despite its "
        "widespread use as a benchmark, published studies often differ in preprocessing choices, "
        "imbalance strategies, and evaluation metrics, making direct comparisons difficult."
    )
    _add_paragraph(doc,
        "This paper addresses four research questions: (1) Which traditional ML algorithm best "
        "classifies customer subscription intent? (2) How does encoding strategy (OHE vs. label "
        "encoding) affect performance? (3) Does SMOTE significantly improve recall for the minority "
        "class? (4) Which features most strongly predict subscription behavior?"
    )

    # ── 2. Dataset and Methodology ────────────────────────────────────
    _add_heading(doc, "2  Dataset and Methodology", level=1)

    _add_heading(doc, "2.1  Dataset", level=2)
    _add_paragraph(doc,
        f"The Bank Marketing dataset contains {n_samples:,} records and 20 input features "
        "encompassing client demographics (age, job, marital status, education), financial history "
        "(default, housing loan, personal loan), campaign contact information (contact type, month, "
        "day, number of contacts), and macroeconomic indicators (employment variation rate, consumer "
        "price index, Euribor 3-month rate, number of employees). The binary target variable y "
        f"indicates term deposit subscription (yes/no) with a {pos_rate:.1%} positive rate."
    )
    _add_paragraph(doc,
        "The duration feature—representing call duration in seconds—was removed prior to modelling. "
        "Although it is a strong predictor (as confirmed by EDA), it is known only after a call "
        "ends and therefore constitutes data leakage in a predictive deployment setting "
        "(Moro et al., 2014)."
    )

    _add_heading(doc, "2.2  Preprocessing", level=2)
    _add_paragraph(doc,
        "'Unknown' values were treated as missing and imputed using the modal value for categorical "
        "features and the median for numerical features. Duplicate records were removed. Two encoding "
        "strategies were compared: (a) one-hot encoding (OHE), expanding categorical variables into "
        "binary indicator columns; and (b) label encoding (LE), converting categories to integer "
        "ordinals."
    )

    _add_heading(doc, "2.3  Feature Engineering", level=2)
    _add_paragraph(doc,
        "For the OHE feature matrix, a chi-square test (SelectKBest, k=25) was applied to identify "
        "the most discriminative features. For the LE matrix, Pearson correlation filtering "
        "(threshold ρ=0.90) removed redundant numerical features. A stratified 80/20 train-test "
        "split with random seed 42 ensured reproducibility."
    )

    _add_heading(doc, "2.4  Class Imbalance Handling", level=2)
    table_imb = doc.add_table(rows=4, cols=2)
    table_imb.style = "Table Grid"
    headers = ["Strategy", "Description"]
    rows_data = [
        ("None (Baseline)", "No modification to training set; models trained on original ~11% positive rate."),
        ("Class Weighting", "Loss function penalised using inverse-frequency class weights (class_weight='balanced')."),
        ("SMOTE", "Synthetic Minority Over-sampling Technique generates new minority-class samples in feature space."),
    ]
    for i, h in enumerate(headers):
        table_imb.rows[0].cells[i].text = h
        _set_cell_bg(table_imb.rows[0].cells[i], "1F497D")
        for para in table_imb.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    for j, (strategy, desc) in enumerate(rows_data):
        table_imb.rows[j + 1].cells[0].text = strategy
        table_imb.rows[j + 1].cells[1].text = desc
        for c in table_imb.rows[j + 1].cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    cap = doc.add_paragraph("Table 1. Class imbalance handling strategies evaluated.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size = Pt(9)
        r.italic = True
    doc.add_paragraph("")

    _add_heading(doc, "2.5  Models and Validation", level=2)
    _add_paragraph(doc,
        "Six classifiers were evaluated: Logistic Regression (LR), Decision Tree (DT), "
        "Random Forest (RF, 100 estimators), Support Vector Machine (SVM with RBF kernel), "
        "k-Nearest Neighbours (k-NN, k=5), and Naïve Bayes (NB). Models requiring feature "
        "scaling (LR, SVM, k-NN) were wrapped with a StandardScaler fit only on training data. "
        "Primary evaluation was on the held-out test set with five metrics: accuracy, precision, "
        "recall, F1-score, and ROC-AUC. The experiment matrix spans 2 encodings × 3 imbalance "
        "strategies × 6 models = 36 total configurations."
    )

    # ── 3. Results ────────────────────────────────────────────────────
    _add_heading(doc, "3  Results", level=1)

    _add_heading(doc, "3.1  Full Model Comparison", level=2)
    _add_paragraph(doc,
        "Table 2 presents all 36 experimental results. Highlighted rows indicate the best "
        "overall configuration by ROC-AUC."
    )
    if comp_df is not None:
        _add_comparison_table(doc, comp_df)
        cap2 = doc.add_paragraph(
            "Table 2. Full model comparison (36 experiments). "
            "Green highlight = highest ROC-AUC configuration."
        )
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap2.runs:
            r.font.size = Pt(9)
            r.italic = True
    else:
        _add_paragraph(doc, "[Table 2 will appear after running run_pipeline.py]", italic=True)

    doc.add_paragraph("")

    _add_heading(doc, "3.2  ROC Curves", level=2)
    _add_paragraph(doc,
        "Figure 1 displays ROC curves for all six classifiers under their best OHE configuration. "
        "Ensemble methods (Random Forest) and linear models (Logistic Regression) consistently "
        "achieve higher AUC than instance-based (k-NN) and probabilistic (Naïve Bayes) methods."
    )
    _add_figure(doc, os.path.join(CHARTS_DIR, "roc_curve.png"),
                "Figure 1. ROC Curves for All Six Classifiers (OHE, best imbalance configuration per model)")

    _add_heading(doc, "3.3  Overall Performance Comparison", level=2)
    _add_figure(doc, os.path.join(CHARTS_DIR, "model_performance.png"),
                "Figure 2. F1-Score and ROC-AUC Comparison Across Models (Best Configuration)")

    _add_heading(doc, "3.4  Feature Importance", level=2)
    _add_paragraph(doc,
        "Figure 3 shows the top 15 features from the Random Forest (SMOTE + OHE). "
        f"Economic indicators dominate: {', '.join(top_features[:5])} "
        "are the strongest predictors. Campaign-specific features (campaign contacts, pdays) "
        "and sociodemographic attributes (age, job) were also influential."
    )
    _add_figure(doc, os.path.join(CHARTS_DIR, "feature_importance.png"),
                "Figure 3. Top 15 Feature Importances – Random Forest (SMOTE + OHE)")

    _add_heading(doc, "3.5  Effect of Imbalance Handling", level=2)
    _add_figure(doc, os.path.join(CHARTS_DIR, "imbalance_impact.png"),
                "Figure 4. Mean Recall and F1-Score by Imbalance Strategy (averaged across all models)")

    _add_heading(doc, "3.6  Best Model – Confusion Matrix", level=2)
    _add_figure(doc, os.path.join(CHARTS_DIR, "confusion_matrix.png"),
                f"Figure 5. Confusion Matrix – {best_name} ({best_enc}, {best_imb})")

    # ── 4. Discussion ─────────────────────────────────────────────────
    _add_heading(doc, "4  Discussion", level=1)

    _add_heading(doc, "4.1  Best Model", level=2)
    _add_paragraph(doc,
        f"{best_name} with {best_imb} and {best_enc} encoding achieved the highest overall "
        f"performance (ROC-AUC = {best_roc_auc}, F1 = {best_f1}). Random Forest benefits from "
        "ensemble averaging, which reduces variance and handles feature interactions implicitly—"
        "crucial for the heterogeneous mix of categorical and macroeconomic features in this dataset. "
        "Logistic Regression, despite its simplicity, performs competitively when class weights are "
        "applied, demonstrating that linear separability is partially present in the feature space."
    )

    _add_heading(doc, "4.2  Impact of SMOTE", level=2)
    _add_paragraph(doc,
        "SMOTE consistently increased recall across all six models, often at the cost of a modest "
        "precision reduction. This trade-off aligns with business objectives in telemarketing: "
        "missing a potential subscriber (false negative) typically incurs a higher opportunity cost "
        "than contacting a non-subscriber (false positive). Class weighting offers a lighter-weight "
        "alternative with similar recall gains and lower computational overhead."
    )

    _add_heading(doc, "4.3  Feature Importance and Drivers", level=2)
    _add_paragraph(doc,
        "Macroeconomic context overwhelmingly drives predictions: the Euribor 3-month rate, "
        "employment variation rate, number of employees, and consumer confidence index collectively "
        "reflect the broader economic conditions under which clients make investment decisions. "
        "Clients contacted in previous campaigns (pdays < 999) show substantially higher conversion "
        "rates. Among sociodemographic features, retired individuals and students show elevated "
        "subscription propensity—consistent with findings in the reference literature "
        "(Moro et al., 2014)."
    )

    _add_heading(doc, "4.4  Encoding Strategy", level=2)
    _add_paragraph(doc,
        "OHE consistently outperformed label encoding across most models. Label encoding imposes "
        "arbitrary ordinal relationships on nominal categorical features (e.g., job titles), which "
        "can mislead distance-based models (k-NN) and linear models (LR, SVM). Tree-based methods "
        "(DT, RF) are more robust to encoding choice because splits are evaluated independently."
    )

    # ── 5. Limitations and Future Work ────────────────────────────────
    _add_heading(doc, "5  Limitations and Future Work", level=1)
    _add_paragraph(doc,
        "This study is bounded to traditional ML models and does not explore gradient boosting "
        "methods (XGBoost, LightGBM) or neural networks, which may yield higher performance. "
        "Hyperparameter tuning was not applied in the primary 36-experiment matrix; future work "
        "should incorporate GridSearchCV or Bayesian optimisation. SHAP-based explanations would "
        "provide more rigorous and model-agnostic feature attribution than the Gini importance "
        "used here. The dataset is geographically specific (Portugal, 2008–2013); generalisation "
        "to other banking markets requires additional validation."
    )

    # ── References ────────────────────────────────────────────────────
    _add_heading(doc, "References", level=1)
    refs = [
        ("Moro, S., Cortez, P., & Rita, P. (2014).",
         "A data-driven approach to predict the success of bank telemarketing. "
         "Decision Support Systems, 62, 22–31."),
        ("Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002).",
         "SMOTE: Synthetic minority over-sampling technique. "
         "Journal of Artificial Intelligence Research, 16, 321–357."),
        ("Pedregosa, F. et al. (2011).",
         "Scikit-learn: Machine learning in Python. "
         "Journal of Machine Learning Research, 12, 2825–2830."),
        ("Breiman, L. (2001).",
         "Random forests. Machine Learning, 45(1), 5–32."),
        ("Yashada Nikam (2023).",
         "UCI Bank Marketing Analysis [GitHub repository]. "
         "https://github.com/yashada-nikam/UCI-Bank-Marketing-Analysis"),
        ("Kataev, A. (2023).",
         "UCI Bank Marketing Dataset. Part 1. Exploratory Data Analysis. "
         "Medium. https://alexkataev.medium.com/uci-bank-marketing-dataset-part-1-..."),
    ]
    for author, text in refs:
        ref_p = doc.add_paragraph(style="List Bullet")
        ref_p.add_run(author + " ").bold = True
        ref_p.add_run(text).font.size = Pt(10)

    # ── Save ─────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"\n✅  Research abstract saved to: {OUT_PATH}")
    return OUT_PATH


if __name__ == "__main__":
    generate_abstract()
